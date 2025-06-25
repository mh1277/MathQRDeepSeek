import os
import time
import requests
from docx import Document
from dotenv import load_dotenv
import sys

# تحميل مفتاح API من ملف .env
load_dotenv()
API_KEY = os.getenv("DEEPSEEK_API_KEY") or "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"
API_URL = "https://openrouter.ai/api/v1/chat/completions"

# الـ Prompt المطلوب
SYSTEM_PROMPT = """
المهمة:
استخرج بدقة متناهية النصوص العربية والمعادلات الرياضية من الصور المقدمة، مع الالتزام بالشروط التالية: مع ترتيب المخرجات وفق ترتيب ترقم الصفحات المرسلة 

متطلبات الإخراج:
1. النصوص العربية:
   - حافظ على التشكيل إن وُجد
   - احترم علامات الترقيم واتجاه النص (يمين لليسار)
   - انقل الهوامش والتعليقات كما تظهر

2. المعادلات الرياضية:
   - استخدم LaTeX نظيف (plain LaTeX)
   - تجنب الحزم غير مدعومة في Word مثل breqn أو physics
   - استخدم الحزم المتوافقة: amsmath, amssymb, amsfonts, IEEEtrantools
   - تأكد من تهيئة الرموز الخاصة:
        - الأقواس: \left( \right) للقياس التلقائي
        - الكسور: \dfrac{}{} للوضوح
        - المتجهات: \vec{} أو \mathbf{}
        - التكاملات/الاشتقاق: \int, \partial، إلخ

3. التنسيق العام:
   - ضع كل معادلة بين \begin{equation} ... \end{equation} للعرض المركزي
   - رقم المعادلات تلقائيًا
   - للجداول الرياضية: استخدم \begin{array}{c|c} ... \end{array}
   - لا تستخدم رموز * أو # عند العناوين

وأريد تدقيق المحتوى العملي على اعتبارك مدرس رياضيات متمرس.
"""

def process_large_docx(input_path, output_path, max_chunk_size=8000):
    """معالجة ملف وورد ضخم على أجزاء"""
    try:
        doc = Document(input_path)
        total_paragraphs = len(doc.paragraphs)
        print(f"تم تحميل الملف: {os.path.basename(input_path)}")
        print(f"عدد الفقرات: {total_paragraphs}")
        
        results = []
        current_chunk = []
        current_size = 0
        
        # تجميع الفقرات في أجزاء
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                text_size = len(text)
                
                # إذا تجاوز الحجم الحد، أرسل الجزء الحالي
                if current_size + text_size > max_chunk_size and current_chunk:
                    chunk_text = "\n".join(current_chunk)
                    processed = process_text_chunk(chunk_text)
                    results.append(processed)
                    current_chunk = []
                    current_size = 0
                    print(f"تم معالجة جزء: {len(results)} - {len(chunk_text)} حرف")
                
                current_chunk.append(text)
                current_size += text_size
            
            # عرض تقدم العمل
            if (i + 1) % 100 == 0 or (i + 1) == total_paragraphs:
                print(f"جاري تجميع الفقرات: {i+1}/{total_paragraphs} ({((i+1)/total_paragraphs)*100:.1f}%)")
        
        # معالجة الجزء الأخير
        if current_chunk:
            chunk_text = "\n".join(current_chunk)
            processed = process_text_chunk(chunk_text)
            results.append(processed)
        
        # حفظ النتائج في ملف وورد جديد
        save_results_to_docx(results, output_path)
        return True
    
    except Exception as e:
        print(f"خطأ في معالجة الملف: {str(e)}")
        return False

def process_text_chunk(text):
    """معالجة جزء من النص باستخدام DeepSeek API"""
    headers = {"Authorization": f"Bearer {API_KEY}"}
    payload = {
        "model": "deepseek/deepseek-chat:free",
        "messages": [
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": f"""قم بتحليل وتدقيق المحتوى التالي مع تطبيق جميع متطلبات الإخراج بدقة:
                
{text}
"""
            }
        ],
        "max_tokens": 8000
    }
    
    try:
        response = requests.post(API_URL, json=payload, headers=headers, timeout=120)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    
    except requests.exceptions.RequestException as e:
        print(f"خطأ في الاتصال: {str(e)}")
        return text  # إرجاع النص الأصلي في حالة الخطأ
    
    except Exception as e:
        print(f"خطأ غير متوقع: {str(e)}")
        return text

def save_results_to_docx(results, output_path):
    """حفظ النتائج في ملف وورد جديد"""
    doc = Document()
    
    # إضافة عنوان المستند
    doc.add_heading('الوثيقة المدققة - إخراج mathword', level=1)
    doc.add_paragraph(f"تاريخ المعالجة: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"عدد الأجزاء: {len(results)}")
    doc.add_paragraph("="*80)
    
    # إضافة المحتوى المدقق
    for i, result in enumerate(results):
        if i > 0:
            doc.add_page_break()
        doc.add_heading(f"الجزء {i+1}", level=2)
        
        # تقسيم النص إلى فقرات مع الحفاظ على التنسيق
        paragraphs = result.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # إذا كانت معادلة
                if "\\begin{equation}" in para:
                    doc.add_paragraph(para)
                else:
                    p = doc.add_paragraph()
                    p.add_run(para)
        
        doc.add_paragraph("-"*80)
    
    # حفظ الملف
    doc.save(output_path)
    print(f"تم حفظ الملف المدقق: {output_path}")

def main():
    """الدالة الرئيسية"""
    print("="*80)
    print("mathword - أداة تدقيق الوثائق الرياضية الاحترافية")
    print("="*80)
    print(f"نظام الـ Prompt المستخدم:\n{'='*40}\n{SYSTEM_PROMPT}\n{'='*40}")
    
    # المسارات الافتراضية
    input_docx = "input.docx"
    output_docx = "output.docx"
    
    # استخدام مسارات من سطر الأوامر إن وجدت
    if len(sys.argv) > 1:
        input_docx = sys.argv[1]
    if len(sys.argv) > 2:
        output_docx = sys.argv[2]
    
    # بدء المعالجة
    start_time = time.time()
    print(f"بدء معالجة الملف: {input_docx}")
    
    success = process_large_docx(input_docx, output_docx)
    
    # إحصائيات الأداء
    elapsed = time.time() - start_time
    mins, secs = divmod(elapsed, 60)
    
    print("\n" + "="*80)
    if success:
        print(f"تمت المعالجة بنجاح في {int(mins)} دقيقة و{int(secs)} ثانية")
        print(f"الملف الناتج: {output_docx}")
    else:
        print("حدث خطأ أثناء المعالجة")
    
    print("="*80)
    input("اضغط Enter للخروج...")

if __name__ == "__main__":
    main()